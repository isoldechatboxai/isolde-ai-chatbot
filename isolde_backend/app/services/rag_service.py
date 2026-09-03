# app/services/rag_service.py
"""
Minimal Retrieval-Augmented-Generation store.

Embeddings and chunks are persisted transactionally in the application
database. PostgreSQL deployments use pgvector cosine distance in the database;
SQLite remains a compatibility-only test backend.
"""

import json
import os
import uuid

import numpy as np

from typing import Optional, List, Dict

from flask import current_app, has_app_context

from app.extensions import db
from app.models.rag_model import RAGChunk, RAGDocument
from app.services.provider_router import embed_text

INDEX_FILENAME = "index.json"
DEFAULT_RELEVANCE_THRESHOLD = 0.55


def _index_path():
    store_dir = current_app.config.get("VECTOR_STORE_DIR", "vector_store")
    os.makedirs(store_dir, exist_ok=True)

    return os.path.join(store_dir, INDEX_FILENAME)


def _load_index(user_id=None):
    """Compatibility projection used by retrieval and legacy callers."""
    query = db.session.query(RAGChunk, RAGDocument.filename).join(
        RAGDocument, RAGDocument.id == RAGChunk.document_id
    )
    if user_id is None:
        query = query.filter(RAGChunk.user_id.is_(None))
    else:
        query = query.filter(RAGChunk.user_id == str(user_id))
    rows = query.all()
    return [{
        "id": chunk.id,
        "file_id": chunk.document_id,
        "filename": filename,
        "user_id": chunk.user_id,
        "text": chunk.text,
        "vector": chunk.embedding,
    } for chunk, filename in rows]


def _uses_postgres():
    # Some legacy unit-level callers exercise ownership filtering without an
    # application context or database. Those callers use the JSON compatibility
    # path; real requests always have an application context.
    return has_app_context() and db.session.get_bind().dialect.name == "postgresql"


def _save_index(records):
    raise RuntimeError("Whole-index replacement is disabled; use transactional RAG operations.")


def _chunk_text(text: str, chunk_size: int = 800, overlap: int = 100):
    chunks = []

    if not text:
        return chunks

    if chunk_size <= 0:
        chunk_size = 800

    if overlap < 0:
        overlap = 0

    if overlap >= chunk_size:
        overlap = max(0, chunk_size // 10)

    start = 0

    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start = end - overlap

    return [chunk.strip() for chunk in chunks if chunk.strip()]


def _extract_pdf_text(file_path: str, filename: str) -> str:
    try:
        import PyPDF2
    except ImportError:
        current_app.logger.error("PyPDF2 is not installed. Run: pip install PyPDF2")
        return ""

    try:
        text_parts = []

        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)

            for page in reader.pages:
                try:
                    extracted = page.extract_text()
                except Exception as page_error:
                    current_app.logger.error(
                        f"Failed to extract a page from {filename}: {page_error}"
                    )
                    continue

                if extracted:
                    text_parts.append(extracted)

        return "\n".join(text_parts)
    except Exception as e:
        current_app.logger.error(f"Failed to extract PDF {filename}: {e}")
        return ""


def _extract_docx_text(file_path: str, filename: str) -> str:
    try:
        import docx
    except ImportError:
        current_app.logger.error(
            "python-docx is not installed. Run: pip install python-docx"
        )
        return ""

    try:
        document = docx.Document(file_path)
        text_parts = []

        for paragraph in document.paragraphs:
            if paragraph.text:
                text_parts.append(paragraph.text)

        for table in document.tables:
            for row in table.rows:
                row_text_parts = []

                for cell in row.cells:
                    if cell.text:
                        row_text_parts.append(cell.text)

                row_text = " | ".join(row_text_parts).strip()

                if row_text:
                    text_parts.append(row_text)

        return "\n".join(text_parts)
    except Exception as e:
        current_app.logger.error(f"Failed to extract DOCX {filename}: {e}")
        return ""


def extract_text_from_file(file_path: str, filename: str) -> str:
    """
    Extract text from supported files.

    Supported:
    - PDF
    - DOCX
    - XLSX
    - TXT
    - CSV
    - Other plain-text files
    """
    if not file_path or not os.path.exists(file_path):
        current_app.logger.error(f"File not found for extraction: {file_path}")
        return ""

    ext = filename.lower().rsplit(".", 1)[-1].strip() if "." in filename else ""

    try:
        if ext == "pdf":
            return _extract_pdf_text(file_path, filename)

        if ext == "docx":
            return _extract_docx_text(file_path, filename)

        if ext == "xlsx":
            try:
                import openpyxl
                wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
                rows_text = []
                for sheet in wb.sheetnames:
                    ws = wb[sheet]
                    rows_text.append(f"[Sheet: {sheet}]")
                    for row in ws.iter_rows(values_only=True):
                        row_values = [str(cell) if cell is not None else "" for cell in row]
                        if any(v.strip() for v in row_values):
                            rows_text.append(" | ".join(row_values))
                wb.close()
                return "\n".join(rows_text)
            except ImportError:
                current_app.logger.error("openpyxl is not installed. Run: pip install openpyxl")
                return f"[XLSX extraction failed. openpyxl not installed. File: {filename}]"
            except Exception as e:
                current_app.logger.error(f"Failed to extract XLSX {filename}: {e}")
                return ""

        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        current_app.logger.error(f"Failed to read file {filename}: {e}")
        return ""


def process_and_index_file(file_path: str, filename: str, user_id: Optional[str] = None):
    """
    Extract text from a file and index it.

    Returns the number of successfully indexed chunks.
    """
    text = extract_text_from_file(file_path, filename)

    if not text or not text.strip():
        return 0

    file_id = str(uuid.uuid4())

    return index_document(file_id, filename, text, user_id)


def index_document(file_id: str, filename: str, text: str, user_id: Optional[str] = None):
    """
    Chunk, embed, and store a document's text for later semantic search.
    """
    if not text or not text.strip():
        return 0

    user_id = str(user_id).strip() if user_id is not None else ""
    if not user_id:
        raise ValueError("Authenticated document ownership is required for RAG indexing.")
    chunks = _chunk_text(text)

    if not chunks:
        return 0

    embedded_chunks = []

    for chunk in chunks:
        try:
            vector = embed_text(chunk)

            if not vector:
                continue

            vector = [float(value) for value in vector]

            embedded_chunks.append((chunk, vector))
        except Exception as e:
            current_app.logger.error(f"Embedding failed for {filename}: {e}")
            continue

    if not embedded_chunks:
        return 0

    document = RAGDocument(id=str(file_id), user_id=user_id, filename=filename, chunk_count=len(embedded_chunks))
    try:
        db.session.add(document)
        for position, (chunk, vector) in enumerate(embedded_chunks):
            db.session.add(RAGChunk(
                document_id=document.id,
                user_id=user_id,
                position=position,
                text=chunk,
                embedding=vector,
                embedding_vector=vector,
                embedding_dimension=len(vector),
            ))
        db.session.commit()
    except Exception:
        db.session.rollback()
        raise
    return len(embedded_chunks)


def list_documents(user_id: str):
    return RAGDocument.query.filter_by(user_id=str(user_id)).order_by(RAGDocument.created_at.desc()).all()


def get_document(file_id: str, user_id: str):
    return RAGDocument.query.filter_by(id=str(file_id), user_id=str(user_id)).first()


def delete_document(file_id: str, user_id: str) -> bool:
    document = RAGDocument.query.filter_by(id=str(file_id), user_id=str(user_id)).first()
    if document is None:
        return False
    try:
        RAGChunk.query.filter_by(document_id=document.id, user_id=str(user_id)).delete(synchronize_session=False)
        db.session.delete(document)
        db.session.commit()
    except Exception:
        db.session.rollback()
        raise
    return True


def import_legacy_json(path: str):
    """Import the former index.json without re-embedding or assigning guessed owners."""
    with open(path, "r", encoding="utf-8") as handle:
        records = json.load(handle)
    if not isinstance(records, list):
        raise ValueError("Legacy RAG index must contain a JSON list.")
    grouped = {}
    skipped_ownerless = 0
    for record in records:
        if not isinstance(record, dict) or not record.get("file_id") or not record.get("text") or not isinstance(record.get("vector"), list):
            raise ValueError("Legacy RAG index contains an invalid record.")
        owner = str(record.get("user_id") or "").strip()
        if not owner:
            skipped_ownerless += 1
            continue
        key = (owner, str(record["file_id"]), str(record.get("filename") or "document"))
        grouped.setdefault(key, []).append(record)
    imported_documents = imported_chunks = 0
    try:
        for (owner, document_id, filename), chunks in grouped.items():
            if RAGDocument.query.filter_by(id=document_id).first():
                continue
            document = RAGDocument(id=document_id, user_id=owner, filename=filename, chunk_count=len(chunks))
            db.session.add(document)
            for position, record in enumerate(chunks):
                vector = [float(value) for value in record["vector"]]
                if not vector:
                    raise ValueError("Legacy RAG embedding cannot be empty.")
                db.session.add(RAGChunk(
                    id=str(record.get("id") or uuid.uuid4()), document_id=document_id,
                    user_id=owner, position=position, text=str(record["text"]),
                    embedding=vector, embedding_dimension=len(vector),
                ))
                imported_chunks += 1
            imported_documents += 1
        db.session.commit()
    except Exception:
        db.session.rollback()
        raise
    return {"documents": imported_documents, "chunks": imported_chunks, "ownerless_skipped": skipped_ownerless}


def search(query: str, top_k: int = 4, user_id: Optional[str] = None) -> List[Dict]:
    """
    Return the top_k most relevant chunks for a query, deduplicated.
    """
    if not query or not query.strip():
        return []

    if top_k <= 0:
        return []

    user_id = str(user_id).strip() if user_id is not None else ""
    if not user_id:
        return []

    try:
        query_vector_raw = embed_text(query)
        query_vec = np.array(query_vector_raw, dtype=float).flatten()
    except Exception as e:
        current_app.logger.error(f"Query embedding failed: {e}")
        return []

    if query_vec.size == 0:
        return []

    # Production PostgreSQL retrieval never materializes every embedding in
    # Python. pgvector's cosine operator is evaluated after owner and
    # dimension filters, preserving isolation and avoiding mixed dimensions.
    if _uses_postgres():
        return _search_postgres(query_vec.tolist(), top_k, user_id)

    records = _load_index(user_id=user_id)

    if not records:
        return []

    scored = []
    seen_texts = set()

    for record in records:
        record_user_id = record.get("user_id")

        if record_user_id is not None:
            record_user_id = str(record_user_id).strip() or None

        # Records are private to their exact authenticated owner. Unowned
        # legacy records are never eligible for retrieval.
        if record_user_id != user_id:
            continue

        chunk_key = record.get("text", "")[:200]
        if chunk_key in seen_texts:
            continue

        try:
            vec = np.array(record.get("vector"), dtype=float).flatten()
        except Exception:
            continue

        if vec.size == 0:
            continue

        if vec.shape != query_vec.shape:
            continue

        denominator = (np.linalg.norm(query_vec) * np.linalg.norm(vec)) or 1e-9
        similarity = float(np.dot(query_vec, vec) / denominator)

        if similarity > DEFAULT_RELEVANCE_THRESHOLD:
            scored.append((similarity, record))
            seen_texts.add(chunk_key)

    scored.sort(key=lambda item: item[0], reverse=True)

    return [
        {
            "text": record["text"],
            "filename": record["filename"],
            "score": round(score, 3),
        }
        for score, record in scored[:top_k]
    ]


def _search_postgres(query_vector: list[float], top_k: int, user_id: str) -> List[Dict]:
    """Exact database-side pgvector cosine retrieval for PostgreSQL only.

    The embedding model is configurable, so no fixed-dimension HNSW/IVFFlat
    index is created. The tenant/dimension B-tree index narrows candidates;
    pgvector then computes exact cosine distance without a Python full scan.
    """
    dimension = len(query_vector)
    distance = RAGChunk.embedding_vector.cosine_distance(query_vector)
    max_distance = 1 - float(current_app.config.get("RAG_RELEVANCE_THRESHOLD", DEFAULT_RELEVANCE_THRESHOLD))
    rows = (
        db.session.query(RAGChunk, RAGDocument.filename, distance.label("distance"))
        .join(RAGDocument, RAGDocument.id == RAGChunk.document_id)
        .filter(
            RAGChunk.user_id == user_id,
            RAGChunk.embedding_dimension == dimension,
            RAGChunk.embedding_vector.isnot(None),
            distance <= max_distance,
        )
        .order_by(distance.asc(), RAGChunk.id.asc())
        .limit(top_k)
        .all()
    )
    return [
        {"text": chunk.text, "filename": filename, "score": round(1 - float(value), 3)}
        for chunk, filename, value in rows
    ]


def build_context_block(chunks: List[Dict], max_chars: int = 4000) -> str:
    """Build context block with a hard character limit to respect token budgets."""
    if not chunks:
        return ""

    lines = ["Relevant knowledge base excerpts:"]
    total_length = len(lines[0])

    for chunk in chunks:
        entry = f"- (from {chunk['filename']}, relevance {chunk['score']}): {chunk['text']}"
        if total_length + len(entry) > max_chars:
            break
        lines.append(entry)
        total_length += len(entry)

    return "\n".join(lines) if len(lines) > 1 else ""
